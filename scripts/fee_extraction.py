"""
Shared, format-agnostic helpers for extracting procedure-code -> fee mappings
out of PT/association fee guides, whatever format they happen to be in
(xlsx, xlsm, xls, csv, docx, or pdf).

The core idea: fee guide tables differ wildly in column layout from one
province/specialty to the next, but they share one property we can exploit --
each data row has a cell holding a recognizable procedure code and another
cell holding a recognizable fee. So instead of hardcoding column positions,
every loader here reduces its source to rows of raw cell values and a single
generic scanner (`extract_codes_from_rows`) looks for (code, fee) pairs,
using the set of procedure codes we actually care about (from the matching
CDCP price file) to avoid false positives on category headers/page numbers.
"""

import csv
import re
from pathlib import Path

import openpyxl
import pypdf
import xlrd
from docx import Document

_DOLLAR_RE = re.compile(r"\d[\d,]*(?:\.\d+)?")

# French-Canadian formatted amount within a spreadsheet text cell (space as
# thousands separator, comma as decimal separator, trailing "$" -- e.g.
# "1 509 $ + L" in QC's SP guide). Matched and removed from the cell *before*
# the plain-digit _DOLLAR_RE pass in _fee_candidates below, which would
# otherwise split "1 509" into separate "1" and "509" tokens (fragmenting
# the thousands digit and, worse, silently discarding it as looking like a
# quantity) rather than reading the number as 1509.
_FRENCH_GROUPED_SPACE_RE = re.compile(r"\d{1,3}(?:[\s ]\d{3})+(?:[.,]\d{2})?\s*\$")

# "S.C." ("Service Charge"/"Independent Charge", per NB's DD guide's own
# abbreviations legend) is yet another no-fixed-fee marker, like "I.C." and
# "c.s." -- but with the two letters in the opposite order, so it does NOT
# match the "c.s." alternative above despite meaning the same thing. Without
# its own alternative here, a cell/segment whose only content is "S.C."
# isn't recognized as a marker at all, so has_no_fee_marker (below, and in
# _FEE_TOKEN_TIERS) never fires for it and it just resolves to nothing.
_NO_FEE_MARKER_RE = re.compile(r"^\s*(?:I\.?\s*C\.?|c\.?\s*s\.?|s\.?\s*c\.?)\s*\.?\s*$", re.IGNORECASE)

SPREADSHEET_SUFFIXES = {".xlsx", ".xlsm", ".xls"}
CSV_SUFFIXES = {".csv"}
DOC_SUFFIXES = {".docx"}
PDF_SUFFIXES = {".pdf"}


_ALPHA_CODE_RE = re.compile(r"^[A-Z]\d{3,5}$")


def normalize_code(raw) -> str | None:
    """Normalize a procedure code (int or str, possibly missing leading zeros) to 5 digits.

    Codes are always whole numbers, so a fractional float (e.g. a $221.75 fee
    cell) is rejected rather than silently truncated -- otherwise a fee could
    coincidentally truncate to a real code (221.75 -> 221 -> "00221") and get
    misread as a code cell in an unrelated row.

    A handful of GP/SP codes are alphanumeric (e.g. "P0500") rather than
    purely numeric -- one letter followed by 3-5 digits is also accepted.
    """
    if raw is None:
        return None
    if isinstance(raw, float) and not raw.is_integer():
        return None
    try:
        return f"{int(raw):05d}"
    except (ValueError, TypeError):
        s = str(raw).strip().upper()
        if s.isdigit() and len(s) <= 6:
            return s
        if _ALPHA_CODE_RE.match(s):
            return s
        return None


def extract_max_dollar(value) -> float | None:
    """Extract the largest dollar amount from a fee cell.

    Handles plain numbers, ranges ("$44.66 to $89.32"), and suffixed fees
    ("$56.02 + exp"). Non-numeric fees (e.g. "c.s." / client specific) return None.
    """
    if value is None:
        return None
    if isinstance(value, (int, float)):
        return float(value)
    numbers = []
    for m in _DOLLAR_RE.findall(str(value)):
        try:
            numbers.append(float(m.replace(",", "")))
        except ValueError:
            continue
    return max(numbers) if numbers else None


def _fee_candidates(cell) -> list[float]:
    """Fee candidates for one non-code cell. Real numeric cells (the normal
    case for a spreadsheet fee column) are always trusted. For text cells,
    extract dollar-like numbers but ignore bare 5-digit tokens -- those are
    almost always a *different* procedure code mentioned in a description
    ("...see code 00616 below"), not a fee -- and ignore small whole numbers
    with no decimal point, since those are usually a size/quantity/duration
    mentioned in a description ("Scar Tissue - 1 - 2 cm", "1 unit of time")
    rather than a fee. This matters most for all-text sources like csv/docx,
    where every cell is a string and there's no numeric type to fall back on
    to distinguish a real fee cell from a description cell."""
    if isinstance(cell, (int, float)):
        return [float(cell)]
    if not isinstance(cell, str):
        return []
    # Collected as (position, value) and sorted at the end so a cell with
    # both kinds of amount in it (e.g. a French range "949 $ - 1 369 $")
    # comes back in left-to-right reading order -- callers take the *last*
    # candidate as the fee (see extract_codes_from_rows), so losing that
    # order would silently swap which end of a range wins.
    matches: list[tuple[int, float]] = []
    masked = cell
    for m in _FRENCH_GROUPED_SPACE_RE.finditer(cell):
        parsed = _parse_french_amount(m.group(0))
        if parsed is not None:
            matches.append((m.start(), parsed))
        # Blank out the matched span (same length, so positions of
        # everything else are unaffected) so the plain-digit pass below
        # doesn't also re-match the digits inside it.
        masked = masked[:m.start()] + " " * (m.end() - m.start()) + masked[m.end():]
    for m in _DOLLAR_RE.finditer(masked):
        token = m.group(0)
        digits = token.replace(",", "")
        # Check the 5-digit-code exclusion against the comma-stripped form,
        # not the raw token: _DOLLAR_RE's character class treats a comma as
        # part of the number, so a code-like reference immediately followed
        # by a comma in prose ("...classes 30000, 40000, and 70000.") comes
        # out as "30000," -- which doesn't literally fullmatch \d{5} (it has
        # 6 characters) even though it's the exact same 5-digit number this
        # exclusion exists to catch.
        if re.fullmatch(r"\d{5}", digits):
            continue
        try:
            value = float(digits)
        except ValueError:
            continue
        if "." not in token and value < 10:
            continue
        matches.append((m.start(), value))
    matches.sort(key=lambda t: t[0])
    return [v for _, v in matches]


# Column headers containing one of these words are treated as a "fee"
# column when looking for which cell in a row holds the fee -- this avoids
# being misled by an unrelated numeric column elsewhere in the row (e.g. a
# sequential "ID"/row-number column) when the *real* fee is stored as
# formatted text ("$89.80") rather than a native number, which otherwise
# defeats the "prefer real numeric cells" heuristic below.
_FEE_HEADER_RE = re.compile(r"fee|amount|price|tarif|montant|\bfrom\b|\brate\b|\bcost\b|\bcharge\b", re.IGNORECASE)

# Column headers containing "code" are treated as *the* code column when
# present. Without this, a fee value that happens to zero-pad to a real
# procedure code (e.g. a $1,802 fee coincidentally matching code "01802")
# gets misread as a reference to that other code -- restricting code
# detection to a labeled code column, when one is identifiable, avoids that
# numeric coincidence entirely.
_CODE_HEADER_RE = re.compile(r"\bcode\b", re.IGNORECASE)


def _looks_like_header(row) -> bool:
    """True if `row` looks like a header row: several short, densely-packed
    text labels, no cell that could plausibly be a procedure code (a real
    header row won't have a 5-digit-normalizable cell -- code columns are
    headed by text like "Code"). Requires at least 2 populated cells and
    short cell text, so a single-cell title/banner row (common as row 1 in
    these fee guides, e.g. "NLDHA ... 2026 Fee Guide ...") isn't mistaken
    for a real header just because it's textual and happens to contain a
    word like "Fee" -- that would misdirect the fee-column search entirely.
    """
    cells = [c for c in row if c is not None]
    if len(cells) < 2:
        return False
    if any(normalize_code(c) for c in cells):
        return False
    if any(isinstance(c, str) and len(c) > 60 for c in cells):
        return False
    if sum(isinstance(c, str) for c in cells) / len(cells) <= 0.5:
        return False
    return len(cells) / len(row) >= 0.3


_EXACT_FEE_HEADER_NAMES = {"fee", "price", "amount", "tarif", "montant", "rate", "cost", "from"}


def find_fee_column_indices(header) -> list[int]:
    """Prefer a column whose header is *exactly* a fee-like word (e.g. "Fee")
    over one that merely contains one as a substring (e.g. "UpperFee",
    "InternalLabFee"). Some guides have both a base fee column and a
    secondary upper-bound-of-range column ("Fee" / "UpperFee") -- the base
    "Fee" column is what the reference file treats as the canonical fee, so
    an exact match should win outright rather than being merged in with (and
    potentially outranked by, since the rightmost match wins) the range
    column.

    A range column pair headed exactly "From"/"To" is the opposite case,
    though (seen in one ON GP sheet): the reference treats the ceiling
    ("To") as the fee, not the floor ("From"). "To" isn't in
    _EXACT_FEE_HEADER_NAMES on its own -- alone it's too generic a word and
    wrongly outranks a real fuzzy-matched fee column in other sheets (one ON
    OS sheet has "Suggested Fee " / "To", where "To" is just an unused
    vestigial column and the real data lives in "Suggested Fee", which would
    get wrongly excluded if "To" won as an exact match by itself) -- so it's
    only added here, and only when "From" is also an exact match in the same
    header, confirming a genuine paired range rather than a stray "To"."""
    exact = [i for i, h in enumerate(header) if h and str(h).strip().lower() in _EXACT_FEE_HEADER_NAMES]
    if exact and any(header[i] and str(header[i]).strip().lower() == "from" for i in exact):
        exact += [i for i, h in enumerate(header) if h and str(h).strip().lower() == "to"]
    if exact:
        return exact
    return [i for i, h in enumerate(header) if h and _FEE_HEADER_RE.search(str(h))]


def find_code_column_indices(header) -> list[int]:
    return [i for i, h in enumerate(header) if h and _CODE_HEADER_RE.search(str(h))]


def extract_codes_from_rows(
    tables, known_codes: set[str], target_specialty: str | None = None
) -> dict[str, float]:
    """Generic (code, fee) scanner for row-shaped tabular data.

    `tables` is an iterable of (header, data_rows) pairs -- one per
    sheet/table, header may be None if none was detected. For each data row,
    any cell that normalizes to a code in `known_codes` is treated as a
    procedure-code cell. If the table has an identifiable "fee" column (by
    header name), only that column is considered -- this is what keeps an
    unrelated numeric column (e.g. a sequential ID) from being mistaken for
    the fee when the real fee is stored as text. Otherwise, the fee is taken
    from actual numeric-typed cells in the row if there are any (the
    reliable case for real spreadsheets), else from a dollar-like number
    found in the other cells' text (for all-string sources like csv/docx).
    In all cases we take the *last* (rightmost) candidate rather than the
    largest: fee tables commonly have the fee as their rightmost relevant
    column, sometimes preceded by unrelated numeric columns (e.g. a
    multi-year price-escalation table where earlier columns hold smaller
    prior-year and unrounded intermediate values) where the largest number
    is not necessarily the current fee.

    `target_specialty`, if given (a CDCP sub-specialty code like "EN"), asks
    the scanner to prefer a row whose own specialty-labeled column (see
    find_row_specialty_column) matches that sub-specialty when the same code
    appears more than once in the source under different specialties (e.g.
    PE's combined GP+SP guide lists code 25781 once under "GP" at $101 and
    again under "END" at $139.30) -- without it, the first row encountered
    always wins regardless of which specialty it actually belongs to.
    """
    fees: dict[str, float] = {}
    # code -> [(specialty_label_or_None, fee), ...], only populated when
    # target_specialty is given, since the single-pass "first row wins" path
    # above is enough (and already validated) for every other caller.
    labeled_candidates: dict[str, list[tuple[str | None, float]]] = {}

    for header, rows in tables:
        fee_col_indices = find_fee_column_indices(header) if header else []
        code_col_indices = find_code_column_indices(header) if header else []
        candidate_code_col_indices = code_col_indices or None  # None = scan every column
        specialty_col_idx = find_row_specialty_column(rows) if target_specialty is not None else None
        for row in rows:
            cells = list(row)
            code_cells = []
            for i, cell in enumerate(cells):
                if candidate_code_col_indices is not None and i not in candidate_code_col_indices:
                    continue
                code = normalize_code(cell)
                if code and code in known_codes:
                    code_cells.append((i, code))
            if not code_cells:
                continue
            # Exclude every cell that restates *this same* code elsewhere in
            # the row (e.g. a truncated numeric id "1011" alongside the full
            # code "01011") -- but NOT cells that happen to match a
            # *different* known code, since that's usually just this row's
            # fee value coincidentally zero-padding to resemble some other,
            # unrelated procedure code (e.g. a $2,116 fee for code 75303
            # looks like code "02116" if that also happens to be a real
            # code) rather than an actual reference to it.
            for i, code in code_cells:
                if target_specialty is None and code in fees:
                    continue
                exclude_indices = {j for j, c in code_cells if c == code} | {i}
                if fee_col_indices:
                    other_cells = [cells[j] for j in fee_col_indices
                                   if j not in exclude_indices and j < len(cells)]
                else:
                    other_cells = [cell for j, cell in enumerate(cells) if j not in exclude_indices]

                fee = None
                numeric_candidates = [float(c) for c in other_cells if isinstance(c, (int, float))]
                # When no fee column is identifiable (header=None, so every
                # column got scanned), a row whose intended fee cell is a
                # standalone "I.C."/"c.s." marker (no fee, by design) can
                # still have some *other*, unrelated real numeric cell in
                # the row (e.g. a page number column) that would otherwise
                # get mistaken for the fee -- skip the numeric fallback
                # entirely in that case rather than risk picking it.
                has_no_fee_marker = not fee_col_indices and any(
                    isinstance(c, str) and _NO_FEE_MARKER_RE.match(c) for c in other_cells
                )
                if has_no_fee_marker:
                    # This row's fee is explicitly non-numeric by design;
                    # skip both the numeric and text fallbacks below
                    # entirely rather than let some *other*, unrelated
                    # numeric cell in the row (e.g. a page number column,
                    # picked up by either fallback) be mistaken for it.
                    continue
                if numeric_candidates:
                    fee = numeric_candidates[-1]
                else:
                    text_candidates = [v for cell in other_cells for v in _fee_candidates(cell)]
                    if text_candidates:
                        fee = text_candidates[-1]
                if fee is None:
                    continue
                if target_specialty is None:
                    fees[code] = fee
                else:
                    label = (_classify_specialty_cell(cells[specialty_col_idx])
                              if specialty_col_idx is not None and specialty_col_idx < len(cells) else None)
                    labeled_candidates.setdefault(code, []).append((label, fee))

    if target_specialty is not None:
        for code, candidates in labeled_candidates.items():
            exact = [f for label, f in candidates if label == target_specialty]
            neutral = [f for label, f in candidates if label is None]
            # First occurrence wins here, not last: unlike the generic
            # single-fee scanner above (where a later column is usually the
            # more current one), a code repeating multiple times *within*
            # one already specialty-scoped source is typically a summary/
            # total row followed by itemized sub-step breakdowns (seen in
            # QC's Endodontie sheet: code 33115's first row is the overall
            # retreatment fee, followed by four rows breaking it into
            # pulpectomy/cleaning/obturation/removal sub-fees) -- the
            # reference treats that first, overall row as the code's fee.
            if exact:
                fees[code] = exact[0]
            elif neutral:
                fees[code] = neutral[0]
            # Else: every row found for this code is labeled for some
            # *other*, specific specialty/context (e.g. only a "GP" and an
            # "LTC" row exist, but the code was asked for as "PA") -- rather
            # than guess by picking whichever happened to come last, leave
            # it unmatched. This is the same situation load_pt_fees_by_
            # subspecialty's GP-fallback is meant to handle (deliberately
            # triggered only when a whole sub-specialty gets zero matches,
            # not per missing code -- see its docstring), so silently
            # substituting a wrong-context value here would both produce an
            # incorrect fee and mask that fallback from ever running.
    return fees


def tables_from_spreadsheet_by_sheet(path: Path):
    """Like tables_from_spreadsheet, but yields (sheet_title, header,
    data_rows) triples -- used when a source splits sub-specialties across
    separate *worksheets* within one workbook, rather than separate files
    (see discover_pt_files) or a per-row specialty column (see
    find_row_specialty_column) -- e.g. QC's combined SP guide, with sheets
    named "Endodontie", "Parodontie", etc."""
    suffix = path.suffix.lower()
    if suffix in {".xlsx", ".xlsm"}:
        wb = openpyxl.load_workbook(path, data_only=True)
        for ws in wb.worksheets:
            rows_iter = ws.iter_rows(values_only=True)
            first = next(rows_iter, None)
            if first is None:
                continue
            if _looks_like_header(first):
                yield ws.title, first, list(rows_iter)
            else:
                yield ws.title, None, [first, *rows_iter]
    elif suffix == ".xls":
        wb = xlrd.open_workbook(str(path))
        for sheet in wb.sheets():
            if sheet.nrows == 0:
                continue
            all_rows = [sheet.row_values(i) for i in range(sheet.nrows)]
            if _looks_like_header(all_rows[0]):
                yield sheet.name, all_rows[0], all_rows[1:]
            else:
                yield sheet.name, None, all_rows
    else:
        raise ValueError(f"Unsupported spreadsheet suffix: {suffix}")


# Substrings (rather than classify_file_specialty's whole-word regex) for
# recognizing a CDCP sub-specialty from a worksheet *title* -- these guides
# use French terms not in SUBSPECIALTY_FILE_MARKERS, and one of QC's sheet
# titles ("M?d - Path - Rad Buccale") has a corrupted accented character
# that breaks clean word-boundary matching, so "DIATRIQUE" (surviving
# fragment of "P?diatrique") and the "PATH"+"RAD" combination handled below
# are chosen specifically to survive that corruption.
_SHEET_SPECIALTY_SUBSTRINGS: dict[str, list[str]] = {
    "EN": ["ENDODONTIE"],
    "OS": ["MAXILLO-FACIALE", "MAXILLOFACIAL"],
    "PA": ["PARODONTIE"],
    "PE": ["DIATRIQUE"],
    "PR": ["PROSTHODONTIE"],
    "OR": ["ORTHODONTIE"],
}


def classify_sheet_specialty(title: str) -> set[str]:
    """Which CDCP sub-specialty code(s), if any, a worksheet title
    identifies -- see tables_from_spreadsheet_by_sheet and
    _SHEET_SPECIALTY_SUBSTRINGS. Empty set means the sheet isn't specific to
    one sub-specialty (e.g. QC's "Section Commune" general-services sheet,
    or a "Membres-*" member-directory sheet, which has no fee data at all)."""
    name = title.upper()
    matches = {code for code, substrings in _SHEET_SPECIALTY_SUBSTRINGS.items()
               if any(s in name for s in substrings)}
    # "M?d - Path - Rad Buccale" combines Oral Medicine, Oral Pathology, and
    # Radiology into one sheet -- all three sub-specialties draw their fee
    # from it.
    if "PATH" in name and "RAD" in name:
        matches |= {"OM", "OP", "RA"}
    return matches


def tables_from_spreadsheet(path: Path):
    """Yields (header, data_rows) per worksheet. The first row of each sheet
    is treated as the header if it looks like one (see _looks_like_header);
    otherwise every row in that sheet is treated as data with no header."""
    suffix = path.suffix.lower()
    if suffix in {".xlsx", ".xlsm"}:
        wb = openpyxl.load_workbook(path, data_only=True)
        for ws in wb.worksheets:
            rows_iter = ws.iter_rows(values_only=True)
            first = next(rows_iter, None)
            if first is None:
                continue
            if _looks_like_header(first):
                yield first, list(rows_iter)
            else:
                yield None, [first, *rows_iter]
    elif suffix == ".xls":
        wb = xlrd.open_workbook(str(path))
        for sheet in wb.sheets():
            if sheet.nrows == 0:
                continue
            all_rows = [sheet.row_values(i) for i in range(sheet.nrows)]
            if _looks_like_header(all_rows[0]):
                yield all_rows[0], all_rows[1:]
            else:
                yield None, all_rows
    else:
        raise ValueError(f"Unsupported spreadsheet suffix: {suffix}")


def tables_from_csv(path: Path):
    with open(path, newline="", encoding="utf-8-sig", errors="replace") as f:
        all_rows = list(csv.reader(f))
    if not all_rows:
        return
    if _looks_like_header(all_rows[0]):
        yield all_rows[0], all_rows[1:]
    else:
        yield None, all_rows


def tables_from_docx(path: Path):
    doc = Document(str(path))
    for table in doc.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if not rows:
            continue
        if _looks_like_header(rows[0]):
            yield rows[0], rows[1:]
        else:
            yield None, rows


def docx_paragraph_text(path: Path) -> str:
    doc = Document(str(path))
    return "\n".join(p.text for p in doc.paragraphs)


def _parse_french_amount(text: str) -> float | None:
    """Parse a French-Canadian formatted amount like "1 026 $" or
    "12 345,67 $" -- a space (or non-breaking space) as the thousands
    separator, a comma as the decimal separator, "$" trailing rather than
    leading. Common throughout Quebec's PDF fee guides. Plain digit-based
    parsing (extract_max_dollar) would otherwise fragment "1 026" into two
    separate numbers ("1" and "026") and grab the wrong one."""
    t = re.sub(r"\s+", "", text).replace("$", "").strip()
    t = t.replace(",", ".")
    try:
        return float(t)
    except ValueError:
        return None


_FRENCH_NUMBER_RE = re.compile(r"\d[\d\s ]*(?:[.,]\d{2})?(?=\s*\$)")


def _parse_french_range_amount(text: str) -> float | None:
    """Parse a French-Canadian amount or range where "$" trails each number
    but there's no thousands-grouping or decimals to distinguish it from a
    plain integer ("97 $", "97 $ - 135 $"). Returns the largest value found,
    consistent with how ranges are handled elsewhere (see extract_max_dollar)."""
    values = [_parse_french_amount(m) for m in _FRENCH_NUMBER_RE.findall(text)]
    values = [v for v in values if v is not None]
    return max(values) if values else None


# Fee-shaped tokens, tried in order of specificity, each paired with the
# parser that turns its matched text into a float:
#   - French-Canadian amounts (space-grouped thousands, comma decimals,
#     trailing "$") -- must come first, before the plain-digit tiers below
#     would otherwise fragment a grouped number like "1 026" in two.
#   - a dollar amount with cents (optionally a range or "+ lab/exp" suffix)
#   - a bare two-decimal amount (some guides drop the "$")
#   - a dollar amount with NO cents ("$109") -- some guides (e.g. Yukon's)
#     quote whole-dollar fees with no decimal point. Without this tier these
#     fall all the way through to the bare-number fallback below, where a
#     later, unrelated number in the description ("...up to the age of 3
#     years...") can outrank the real fee.
#   - a "client specific" / "c.s." marker (no fixed fee -- extract_max_dollar
#     finds no digits in it and returns None, which is the correct result)
#   - finally a bare whole number with no "$" at all (some guides, e.g.
#     Quebec's, list fees with no dollar sign or decimals -- riskiest, so
#     tried last and only within a code's own text segment, see
#     extract_codes_from_text).
_FEE_TOKEN_TIERS = [
    (re.compile(r"\d{1,3}(?:[\s ]\d{3})+(?:[.,]\d{2})?\s*\$"), _parse_french_amount),
    (re.compile(r"\d+[.,]\d{2}\s*\$"), _parse_french_amount),
    # Bare digits with trailing "$" and no thousands-grouping or decimals
    # ("97 $", "97 $ - 135 $") -- the two tiers above require grouping or a
    # decimal to distinguish a French-formatted amount from an unrelated
    # bare number, but not every French guide's fees are large/precise
    # enough to have either.
    (re.compile(r"\d+\s*\$(?:\s*(?:to|-)\s*\d+\s*\$)?"), _parse_french_range_amount),
    (re.compile(r"\$[\d,]+\.\d{2}(?:\s*(?:to|-)\s*\$?[\d,]+\.\d{2})?(?:\s*\+\s*[A-Za-z]+)*"), extract_max_dollar),
    (re.compile(r"\b[\d,]+\.\d{2}\b(?:\s*(?:to|-)\s*\$?[\d,]+\.\d{2})?(?:\s*\+\s*[A-Za-z]+)*"), extract_max_dollar),
    (re.compile(r"\$[\d,]+(?:\s*(?:to|-)\s*\$?[\d,]+)?(?:\s*\+\s*[A-Za-z]+)*"), extract_max_dollar),
    (re.compile(r"c\.?\s*s\.?\s*\(?client specific\)?\.?", re.IGNORECASE), extract_max_dollar),
    (re.compile(r"c\.?\s*s\.?", re.IGNORECASE), extract_max_dollar),
    # "I.C." ("Individually Costed") -- another no-fixed-fee marker, same
    # idea as "c.s." above. Without recognizing it, a code's segment like
    # "74112 1 - 2 cm I.C." falls through every tier here to the risky
    # bare-whole-number fallback below, which then misreads the "2" from
    # the size range ("1 - 2 cm") in the description as if it were a fee.
    (re.compile(r"\bI\.?\s*C\.?\b", re.IGNORECASE), extract_max_dollar),
    # "S.C." ("Service Charge"/"Independent Charge", per NB's DD guide) --
    # the same "no fixed fee" idea as "c.s." above, but with the two letters
    # reversed, so the "c.s." tier never matches it. Seen 136 times in NB's
    # DD guide alone; without this tier, every one of those rows' segments
    # (no digits at all -- just the word "S.C." itself) fails to match
    # anything all the way down to the bare-number tier too, so the code
    # never resolves at all, rather than correctly resolving to "no fee".
    (re.compile(r"\bS\.?\s*C\.?\b", re.IGNORECASE), extract_max_dollar),
]
_BARE_NUMBER_TIER = (re.compile(r"\b\d{1,4}\b"), extract_max_dollar)

_ANY_CODE_RE = re.compile(r"\b\d{5}\b")

# How far into a code's text segment to look for its fee. Keeps the
# whole-number fallback tier from wandering into an unrelated number many
# sentences later.
_SEGMENT_SEARCH_WINDOW = 700


_YEAR_RE = re.compile(r"^(19|20)\d{2}$")


def _fee_token_in_segment(segment: str) -> tuple[float | None, bool]:
    """Find the fee within one code's text segment.

    For the specific tiers (dollar amounts, decimals, "c.s."/"I.C." markers)
    the fee is the *first* match -- these guides put the fee right after the
    code ("00111 $63.80..."), or it's the only such token in a descriptive
    paragraph. The bare-whole-number fallback tier is different: guides that
    use it often prefix the fee with a quantity label ("00211 1 image 46"),
    so there the fee is the *last* match instead.

    Returns (fee, found_something) -- `found_something` is True whenever any
    tier matched at all, even a marker tier that parses to no numeric fee
    (e.g. "I.C."/"c.s."), and False only when nothing resembling a fee token
    was found anywhere in the window. The caller (extract_codes_from_text)
    uses this to decide whether this occurrence of the code is authoritative
    -- a marker match means "yes, this is the code's real entry, and it
    genuinely has no fixed fee," which is different from a segment that's
    just descriptive prose mentioning the code in passing (e.g. an intro
    paragraph referencing it before its actual price-table entry appears
    later) and shouldn't block a later, real entry from being found.
    """
    window = segment[:_SEGMENT_SEARCH_WINDOW]
    for pattern, parser in _FEE_TOKEN_TIERS:
        last_valid = None
        for match in pattern.finditer(window):
            text = match.group(0)
            if _YEAR_RE.match(text.strip()):
                continue
            last_valid = text
        if last_valid is not None:
            # If a tier matches more than once in the window (e.g. a
            # "CODE PROF LAB TOTAL" row rendered as plain tab-separated text
            # "31112  1118.00  746.00  1864.00"), the last match is the
            # running total, which is what the reference treats as the
            # code's fee -- and it's harmless when there's only one match,
            # since first and last are then the same token.
            return parser(last_valid), True

    bare_pattern, bare_parser = _BARE_NUMBER_TIER
    last_valid = None
    for match in bare_pattern.finditer(window):
        text = match.group(0)
        if _YEAR_RE.match(text.strip()):
            continue
        last_valid = text
    if last_valid is not None:
        return bare_parser(last_valid), True
    return None, False


# Words that typically precede a *cross-reference* to another code inside a
# description ("...as per 00100", "...see code 00616", "...use code 00616")
# rather than that code's own entry. Occurrences preceded by one of these are
# ignored entirely -- both as segment boundaries and as extraction points --
# so a mention of another code doesn't truncate the current entry's segment,
# and a mention of *this* code doesn't get mistaken for its own definition.
_CROSS_REF_CUE_WORDS = {
    "per", "code", "codes", "see", "use", "using", "refer",
    "of", "or", "and", "lieu", "than", "instead", "not",
    "et",  # French "and" -- QC guides list codes together in French prose
}
_PRECEDING_WORD_RE = re.compile(r"([A-Za-z]+)\s*$")
_FOLLOWING_WORD_RE = re.compile(r"^\s*([A-Za-z]+)")
_FOLLOWED_BY_PERIOD_RE = re.compile(r"^\s?\.")
_FOLLOWED_BY_RANGE_RE = re.compile(r"^\s*-")
_FOLLOWED_BY_LIST_PUNCTUATION_RE = re.compile(r"^\s*[,)]")


def _is_cross_reference(text: str, pos: int) -> bool:
    preceding = _PRECEDING_WORD_RE.search(text[max(0, pos - 30):pos])
    if preceding and preceding.group(1).lower() in _CROSS_REF_CUE_WORDS:
        return True
    end = pos + 5
    # A code immediately followed by a cue word ("...01200 et 01250)") is
    # sitting in the *middle* of a list of codes mentioned together in
    # prose -- not preceded by a cue word itself (that catches the last
    # item, "01250"), and not followed by list punctuation either (that
    # catches earlier items like "01120,"), so it needs its own check.
    following_word = _FOLLOWING_WORD_RE.match(text[end:end + 10])
    if following_word and following_word.group(1).lower() in _CROSS_REF_CUE_WORDS:
        return True
    following = text[end:end + 2]
    # A code immediately followed by a sentence-ending period ("...listed in
    # 00100.") reads as a citation closing out someone else's sentence, not
    # the start of this code's own entry (which is normally followed by more
    # descriptive text, not punctuation).
    if _FOLLOWED_BY_PERIOD_RE.match(following):
        return True
    # A code immediately followed by "-<digit>" ("04401-02") is a code-range
    # reference, typically from an alphabetical index/appendix section
    # ("Dental Legal Letters, 93121-23") rather than the code's own fee
    # entry -- a real entry is followed by a space then descriptive text.
    if _FOLLOWED_BY_RANGE_RE.match(following):
        return True
    # A code immediately followed by "," or ")" is part of a comma-separated
    # or parenthetical list of codes mentioned together in prose (e.g. a
    # French explanatory note "01120, 01130, 01200 et 01250"), not its own
    # entry.
    return bool(_FOLLOWED_BY_LIST_PUNCTUATION_RE.match(following))


def extract_codes_from_text(text: str, known_codes: set[str]) -> dict[str, float]:
    """Segment `text` by occurrences of *any* 5-digit code (not just ones we
    care about), then look for a fee token within each known code's segment
    (the text up to the next code of any kind). Bounding on any code --
    rather than only known ones -- keeps an unrelated nearby entry's numbers
    (for a code outside our CDCP list) from bleeding into the segment.
    Cross-reference mentions of a code within another entry's description are
    excluded from consideration entirely (see _is_cross_reference).

    The *first* legitimate (non-cross-reference) occurrence of a code wins
    once it actually resolves something -- either a numeric fee, or an
    explicit "no fixed fee" marker like "I.C."/"c.s." (see
    _fee_token_in_segment) -- even when that something is "no fixed fee":
    some combined guides (e.g. a "GP SP Fee Guide" covering both general and
    specialist rates in one PDF) list the same code a second time much later
    under a different, specialist-rate section with a real number, and that
    later occurrence must not be mistaken for this code's definition just
    because the first one's fee is genuinely variable. But a segment with no
    recognizable fee token *at all* (e.g. an intro paragraph mentioning the
    code by name well before its actual price-table entry, with no price
    nearby) isn't treated as authoritative, so a later, real entry still
    gets a chance -- otherwise a code mentioned in passing before its own
    definition would never resolve.
    """
    all_matches = [
        m for m in _ANY_CODE_RE.finditer(text) if not _is_cross_reference(text, m.start())
    ]

    fees: dict[str, float] = {}
    seen: set[str] = set()
    for i, m in enumerate(all_matches):
        code = m.group(0)
        if code not in known_codes or code in seen:
            continue
        next_start = all_matches[i + 1].start() if i + 1 < len(all_matches) else len(text)
        segment = text[m.end():next_start]
        fee, found_something = _fee_token_in_segment(segment)
        if found_something:
            seen.add(code)
        if fee is not None:
            fees[code] = fee
    return fees


def load_fees_from_pdf(path: Path, known_codes: set[str]) -> dict[str, float]:
    reader = pypdf.PdfReader(str(path))
    plain_text = "\n".join(page.extract_text() or "" for page in reader.pages)
    plain_fees = extract_codes_from_text(plain_text, known_codes)
    if len(plain_fees) >= len(known_codes):
        return plain_fees
    # Some guides' fee tables are laid out in columns that pypdf's default
    # (draw-order) extraction serializes column-major -- every description,
    # then every code, then every fee, each in its own block far from the
    # others -- which breaks the code-adjacent-to-its-own-fee assumption
    # extract_codes_from_text relies on entirely (seen in NB's DD guide: only
    # 24/398 codes resolved this way, most to the wrong fee). "layout" mode
    # instead reconstructs each row using the text's actual on-page position,
    # so a row's code and fee end up next to each other again (331/398 for
    # the same file). Only tried as a fallback, and only kept if it actually
    # resolves more codes than the plain pass did -- for a source where plain
    # mode already works (the common case), layout's heavier padding and
    # different line-wrapping isn't worth risking a regression on.
    layout_text = "\n".join(
        page.extract_text(extraction_mode="layout") or "" for page in reader.pages
    )
    layout_fees = extract_codes_from_text(layout_text, known_codes)
    return layout_fees if len(layout_fees) > len(plain_fees) else plain_fees


def load_fees_from_docx(path: Path, known_codes: set[str], target_specialty: str | None = None) -> dict[str, float]:
    fees = extract_codes_from_rows(tables_from_docx(path), known_codes, target_specialty)
    missing = known_codes - fees.keys()
    if missing:
        # Some docx fee guides are prose/paragraphs rather than tables.
        fees.update({
            code: fee
            for code, fee in extract_codes_from_text(docx_paragraph_text(path), known_codes).items()
            if code not in fees
        })
    return fees


def load_fees_from_spreadsheet(path: Path, known_codes: set[str], target_specialty: str | None = None) -> dict[str, float]:
    if target_specialty is not None:
        # If any worksheet's title identifies it as specific to the
        # requested sub-specialty (e.g. QC's combined SP guide splits
        # sub-specialties across sheets like "Endodontie", "Parodontie"),
        # prefer those sheets over the rest of the workbook -- otherwise
        # the same code appearing on multiple specialty sheets (a common
        # diagnostic/radiograph code, say) resolves to whichever sheet
        # happens to come first, regardless of which specialty was asked
        # for.
        titled_tables = list(tables_from_spreadsheet_by_sheet(path))
        if any(classify_sheet_specialty(title) for title, _, _ in titled_tables):
            specific = [(h, r) for title, h, r in titled_tables
                        if target_specialty in classify_sheet_specialty(title)]
            fees = extract_codes_from_rows(specific, known_codes, target_specialty)
            missing = known_codes - fees.keys()
            if missing:
                general = [(h, r) for title, h, r in titled_tables if not classify_sheet_specialty(title)]
                fees.update({c: f for c, f in extract_codes_from_rows(general, missing, target_specialty).items()
                             if c not in fees})
            return fees
    return extract_codes_from_rows(tables_from_spreadsheet(path), known_codes, target_specialty)


def load_fees_from_csv(path: Path, known_codes: set[str], target_specialty: str | None = None) -> dict[str, float]:
    return extract_codes_from_rows(tables_from_csv(path), known_codes, target_specialty)


PROVINCE_ALIASES: dict[str, list[str]] = {
    "PE": ["PE", "PEI"],
    "YT": ["YK", "YT", "YU", "YUKON"],
    "YK": ["YK", "YT", "YU", "YUKON"],
    "NT": ["NT", "NWT"],
}


def discover_pt_files(specialty_dir: Path, province: str) -> list[Path]:
    """Find every file relevant to one province under `specialty_dir`.

    Two ways a file can qualify:
    1. Its name starts with the province's abbreviation (or a known alias),
       wherever it lives (covers most cases, including files inside a
       same-named subfolder like SP/BC/BC PA Fee Guide 2026.xlsx).
    2. It lives inside a per-province subfolder that doesn't itself follow
       the naming convention (e.g. SP/MB/MDA 2026 Endo....xlsx -- Manitoba's
       sub-specialty fee guides are split into several files named after the
       vendor, not the province).
    """
    aliases = PROVINCE_ALIASES.get(province, [province])
    patterns = [re.compile(rf"^{re.escape(a)}\b", re.IGNORECASE) for a in aliases]

    province_subdirs = [
        specialty_dir / alias for alias in aliases if (specialty_dir / alias).is_dir()
    ]

    matches = []
    for path in specialty_dir.rglob("*"):
        if not path.is_file():
            continue
        if any(p.match(path.name) for p in patterns):
            matches.append(path)
        elif any(sub in path.parents for sub in province_subdirs):
            matches.append(path)
    return matches


def _is_english(path: Path) -> bool:
    name = path.stem.upper()
    return "FR" not in name.split() and "FRENCH" not in name and "- FR" not in name.upper()


def load_pt_fees_from_files(
    files: list[Path], known_codes: set[str], verbose: bool = True, target_specialty: str | None = None
):
    """Resolve PT fees for `known_codes` from a specific set of files, in
    priority order (filling in only the codes still missing at each step, so
    multiple partial sources combine): spreadsheets (xlsx/xlsm/xls), csv,
    docx, then pdf (English files before French, if language is discernible).

    `target_specialty`, if given, is passed down to the row-based readers
    (spreadsheet/csv/docx) so a combined guide with its own per-row
    specialty column (see find_row_specialty_column) can prefer the row
    actually labeled for that sub-specialty over whichever row happens to
    come first -- not applied to the pdf reader, which uses a different,
    text-segment-based scanner (see extract_codes_from_text).

    Returns (fees dict, list of (source_description, codes_found) used).
    """
    fees: dict[str, float] = {}
    sources_used: list[tuple[str, int]] = []

    def _apply(label: str, new_fees: dict[str, float]):
        added = {c: f for c, f in new_fees.items() if c not in fees}
        fees.update(added)
        if added:
            sources_used.append((label, len(added)))

    def _is_regional_variant(f: Path) -> bool:
        # Some provinces publish a separate fee guide for remote/northern
        # regions (e.g. "MB GP 2026 Fee Guide - NORTHERN.xlsx" alongside the
        # regular "MB GP 2026 Fee Guide.xlsx") with a premium over standard
        # rates. The reference file treats the standard-rate guide as
        # canonical, so it should be tried before, not after, a regional
        # variant -- otherwise the variant's higher rate wins for any code
        # both files list.
        return "NORTHERN" in f.stem.upper()

    spreadsheets = sorted(
        (f for f in files if f.suffix.lower() in SPREADSHEET_SUFFIXES),
        key=lambda f: (f.suffix.lower() != ".xlsx", _is_regional_variant(f)),  # prefer .xlsx, then standard-rate files
    )
    for f in spreadsheets:
        try:
            _apply(f.name, load_fees_from_spreadsheet(f, known_codes, target_specialty))
        except Exception as e:
            if verbose:
                print(f"    WARNING: failed to read {f.name}: {e}")

    for f in (f for f in files if f.suffix.lower() in CSV_SUFFIXES):
        if known_codes - fees.keys():
            try:
                _apply(f.name, load_fees_from_csv(f, known_codes, target_specialty))
            except Exception as e:
                if verbose:
                    print(f"    WARNING: failed to read {f.name}: {e}")

    for f in (f for f in files if f.suffix.lower() in DOC_SUFFIXES):
        if known_codes - fees.keys():
            try:
                _apply(f.name, load_fees_from_docx(f, known_codes, target_specialty))
            except Exception as e:
                if verbose:
                    print(f"    WARNING: failed to read {f.name}: {e}")

    pdfs = sorted(
        (f for f in files if f.suffix.lower() in PDF_SUFFIXES),
        key=lambda f: not _is_english(f),  # English first
    )
    for f in pdfs:
        if known_codes - fees.keys():
            try:
                _apply(f.name, load_fees_from_pdf(f, known_codes))
            except Exception as e:
                if verbose:
                    print(f"    WARNING: failed to read {f.name}: {e}")

    # Last resort, only for whatever spreadsheet/csv/docx/pdf-text all
    # failed to resolve: some PDFs (e.g. NL's DD guide) have no extractable
    # text at all -- their text was flattened to vector curves on export --
    # so nothing above can ever find anything in them no matter how the
    # content stream is read. OCR-ing the rendered page image is the only
    # way to recover data from a file like that. Deliberately tried last and
    # only for the remaining gap (never re-tried on a pdf that already
    # resolved fine above) since it's slow and occasionally misreads a
    # digit, and only imported here so a machine without the OCR
    # dependencies installed (see ocr_pdf_fees.py) just skips this tier
    # instead of failing the whole extraction run.
    if pdfs and (known_codes - fees.keys()):
        try:
            from ocr_pdf_fees import load_fees_from_pdf_via_ocr
        except ImportError as e:
            if verbose:
                print(f"    WARNING: OCR fallback unavailable ({e}); skipping")
        else:
            for f in pdfs:
                remaining = known_codes - fees.keys()
                if not remaining:
                    break
                try:
                    _apply(f"{f.name} (OCR)", load_fees_from_pdf_via_ocr(f, remaining))
                except Exception as e:
                    if verbose:
                        print(f"    WARNING: OCR failed for {f.name}: {e}")

    return fees, sources_used


def load_pt_fees(specialty_dir: Path, province: str, known_codes: set[str], verbose: bool = True):
    """Resolve PT fees for a province/specialty from whatever files are available.
    Returns (fees dict, list of (source_description, codes_found) used, files found)."""
    files = discover_pt_files(specialty_dir, province)
    fees, sources_used = load_pt_fees_from_files(files, known_codes, verbose)
    return fees, sources_used, files


# DD (denturist) fee guides commonly break a procedure's fee into separate
# Professional / Lab / Total columns (e.g. SK: "Prof Fee"/"Lab Fee"/"Total
# Fee", QC (French): "Honoraires"/"Frais de lab."/"Total", BC/MB/NT/NU/ON:
# "PROF"/"LAB"/"TOTAL" headers repeated before every section of the price
# table rather than once at the top of the sheet). The generic single-fee
# scanner above can't represent three distinct values per code, so DD gets
# its own role-aware column classifier and row scanner.
_DD_PROF_HEADER_RE = re.compile(r"\bprof(essional)?\b|honoraires", re.IGNORECASE)
_DD_LAB_HEADER_RE = re.compile(r"\blab\b|laboratoire|frais\s*de\s*lab", re.IGNORECASE)
_DD_TOTAL_HEADER_RE = re.compile(r"\btotal\b", re.IGNORECASE)
_YEAR_TOKEN_RE = re.compile(r"\b(19\d{2}|20\d{2})\b")
# This project's fee guides are all for the 2026 rate year -- a header
# naming *that* year ("Total Fee 2026") is still the current column, but one
# naming an earlier year ("2025 Prof Fee") is a prior-year column to ignore.
_CURRENT_GUIDE_YEAR = 2026


def _dd_column_role(cell) -> str | None:
    """Classify one header cell as 'prof', 'lab', 'total', or None. A cell
    naming a *prior* year (e.g. "2025 Prof Fee") is excluded from all three
    -- these guides list last year's Prof/Lab/Total fee alongside the
    current one under the same style of label, and only the current-year (or
    undated) column is this year's actual fee."""
    if not isinstance(cell, str):
        return None
    text = cell.strip()
    if not text:
        return None
    if any(int(m.group(0)) < _CURRENT_GUIDE_YEAR for m in _YEAR_TOKEN_RE.finditer(text)):
        return None
    if _DD_TOTAL_HEADER_RE.search(text):
        return "total"
    if _DD_LAB_HEADER_RE.search(text):
        return "lab"
    if _DD_PROF_HEADER_RE.search(text):
        return "prof"
    return None


def _find_dd_role_column_candidates(row) -> dict[str, list[int]]:
    """Every column matching each of Prof/Lab/Total, by header text, in
    left-to-right order -- a role label can legitimately appear more than
    once in the same header row (e.g. one QC sheet lists an unlabeled
    prior-year "Honoraires" column alongside the current one under the same
    text; one NS sheet lists a rounded current-year "TOTAL" alongside an
    unrounded prior-year "TOTAL"). Which occurrence is the real current-year
    one isn't consistently the first or the last across sources, so all
    candidates are kept and extract_dd_codes_from_rows picks between them
    using the actual data (whichever Total most closely equals Prof + Lab)
    rather than guessing from position."""
    candidates: dict[str, list[int]] = {}
    for i, cell in enumerate(row):
        role = _dd_column_role(cell)
        if role:
            candidates.setdefault(role, []).append(i)
    return candidates


def _looks_like_dd_role_header(row) -> bool:
    """True if `row` is itself a Prof/Lab/Total-labeled header row, as
    opposed to a data row. Some DD guides (e.g. BC's, MB's, NT's, NU's,
    ON's) repeat this header before every section of the price table rather
    than listing it once at the top of the sheet -- requiring 2+ identified
    roles avoids a data row that merely mentions "lab" in a description
    cell being mistaken for one."""
    return len(_find_dd_role_column_candidates(row)) >= 2


def extract_dd_codes_from_rows(tables, known_codes: set[str]) -> dict[str, dict[str, float]]:
    """DD-specific (code -> {'prof': .., 'lab': .., 'total': ..}) scanner.
    Like extract_codes_from_rows, but for sources with identifiable
    Prof/Lab/Total columns (see _dd_column_role) instead of one generic fee
    column. The active column mapping is re-detected from any row that
    looks like a role header, not just the table's own detected header,
    since several guides repeat it before every section rather than listing
    it once. It also carries forward across tables that don't have (or
    repeat) a header of their own -- some docx guides (e.g. AB's) are a long
    series of small per-section tables sharing one consistent column layout,
    where only a handful of sections actually repeat the "DAC CODE /
    PROFESSIONAL FEE / LAB FEE / TOTAL FEE" header. Deliberately NOT
    backfilled the other direction (using a mapping to reinterpret tables
    that came *before* it was found): tried that and it broke NB's sheet,
    where a Prof/Lab/Total-labeled section buried in the middle of the file
    (digital denture services) got treated as the layout for unrelated
    earlier sections, producing wrong-but-plausible-looking $0 fees instead
    of correctly leaving them for the single-fee fallback. A code with no
    active role mapping yet (no header seen so far in this file) or no fee
    found in any active role column is simply skipped -- the caller falls
    back to the generic single-fee scanner for those.
    """
    results: dict[str, dict[str, float]] = {}
    active_candidates: dict[str, list[int]] = {}
    code_col_indices: list[int] = []
    for header, rows in tables:
        if header:
            header_candidates = _find_dd_role_column_candidates(header)
            if header_candidates:
                active_candidates = header_candidates
            header_code_cols = find_code_column_indices(header)
            if header_code_cols:
                code_col_indices = header_code_cols
        for row in rows:
            cells = list(row)
            if _looks_like_dd_role_header(cells):
                active_candidates = _find_dd_role_column_candidates(cells)
                detected_code_cols = find_code_column_indices(cells)
                if detected_code_cols:
                    code_col_indices = detected_code_cols
                continue
            if not active_candidates:
                continue
            candidate_code_col_indices = code_col_indices or None
            code = None
            for i, cell in enumerate(cells):
                if candidate_code_col_indices is not None and i not in candidate_code_col_indices:
                    continue
                c = normalize_code(cell)
                if c and c in known_codes:
                    code = c
                    break
            if code is None or code in results:
                continue

            def _role_cells(role: str) -> list[tuple[object, float | None]]:
                out = []
                for idx in active_candidates.get(role, []):
                    if idx >= len(cells):
                        continue
                    cell = cells[idx]
                    out.append((cell, extract_max_dollar(cell)))
                return out

            def _is_blank(cell) -> bool:
                return cell is None or (isinstance(cell, str) and not cell.strip())

            prof_cells = _role_cells("prof")
            lab_cells = _role_cells("lab")
            total_candidates = [f for _, f in _role_cells("total") if f is not None]
            prof_candidates = [f for _, f in prof_cells if f is not None]
            lab_candidates = [f for _, f in lab_cells if f is not None]
            # A genuinely blank Prof or Lab cell (as opposed to a
            # non-numeric marker like "+L"/"SC" for a variable/
            # client-specific charge) means this procedure simply has no
            # component on that side -- 0, not unknown -- matching the
            # reference's convention (and load_cdcp_dd_fees's `or 0` for the
            # CDCP side). Only default one side to 0 this way when the
            # *other* side has a real value, so a row where both Prof and
            # Lab columns are blank (e.g. a flat, undivided fee) isn't
            # forced to a fake $0 + $0 instead of being left for Total alone.
            if not prof_candidates and lab_candidates and prof_cells and _is_blank(prof_cells[-1][0]):
                prof_candidates = [0.0]
            if not lab_candidates and prof_candidates and lab_cells and _is_blank(lab_cells[-1][0]):
                lab_candidates = [0.0]
            # Any of the three roles can end up with more than one matching
            # column: a source can repeat the same "Prof Fee"/"Lab Fee"
            # header for its unlabeled prior-year figure alongside the
            # current one (SK's denture sections), or repeat "Total" for a
            # rounded current-year figure alongside an unrounded prior-year
            # one (NS) -- in either left-right order, so position alone
            # can't tell current from prior. Instead, pick whichever
            # combination of candidates actually satisfies this row's own
            # Prof + Lab = Total, rather than guessing by position.
            prof = prof_candidates[-1] if prof_candidates else None
            lab = lab_candidates[-1] if lab_candidates else None
            total = total_candidates[-1] if total_candidates else None
            if len(prof_candidates) > 1 or len(lab_candidates) > 1 or len(total_candidates) > 1:
                best = None
                best_diff = None
                # Iterate right-to-left so that on an exact tie (seen in one
                # QC sheet, where a uniform escalation factor means BOTH the
                # prior-year and current-year triples satisfy Prof+Lab=Total
                # exactly), the rightmost/current-year combination is found
                # first and a later, equally-good match doesn't overwrite it
                # -- consistent with this file's "rightmost = current"
                # convention (see extract_codes_from_rows, _find_dd_role_
                # column_candidates).
                for p in reversed(prof_candidates or [None]):
                    for l in reversed(lab_candidates or [None]):
                        for t in reversed(total_candidates or [None]):
                            if p is None or l is None or t is None:
                                continue
                            diff = abs((p + l) - t)
                            if best is None or diff < best_diff:
                                best, best_diff = (p, l, t), diff
                if best is not None:
                    prof, lab, total = best

            values: dict[str, float] = {}
            if prof is not None:
                values["prof"] = prof
            if lab is not None:
                values["lab"] = lab
            if total is not None:
                values["total"] = total
            if values:
                results[code] = values
    return results


_DD_PDF_LINE_CODE_RE = re.compile(r"\b(\d{5})\b")
_DD_PDF_NUMBER_RE = re.compile(r"\d[\d,]*\.\d{2}")


def extract_dd_codes_from_pdf_text(text: str, known_codes: set[str]) -> dict[str, dict[str, float]]:
    """DD-specific PDF scanner for guides (so far only NB's) that print a
    code's Prof/Lab/Total fees together on the same line, in that left-to-
    right order -- e.g. "Diagnostic Model - Maxillary   10120   122.00
    184.00   306.00" (confirmed left-to-right by Total == Prof + Lab in
    every such line). `text` should be pdf "layout" extraction-mode text
    (see load_fees_from_pdf) -- plain-mode text serializes multi-column
    tables in draw order, scattering a row's numbers away from its code
    entirely, so there'd be nothing to find on the code's own line at all.

    The generic single-fee PDF scanner (load_fees_from_pdf) can only ever
    take one number per code (the rightmost, i.e. Total), silently losing
    the Prof/Lab breakdown for any such row. This recovers it directly from
    the numbers on the same line as the code, without needing a
    recognizable column header -- NB's own header ("CODE / CLINICAL FEE /
    TOTAL FEE") doesn't even name a "Lab" column at all; the breakdown only
    shows up as a third number on the rows that have one.

    A line with exactly one number after the code is Total only (Prof/Lab
    genuinely undifferentiated, same as the generic scanner). Exactly two
    IDENTICAL numbers is common for procedures with no lab component at all
    (Clinical Fee == Total Fee); recorded as Total with Lab forced to 0.0,
    not left unknown, matching this project's established "a genuinely
    blank side is 0, not unknown" DD convention (see
    extract_dd_codes_from_rows). Three or more numbers is read as (prof,
    lab, total) -- but only kept as such if it actually satisfies prof +
    lab == total (within a cent); otherwise treated as Total-only, so a
    line with other, unrelated numbers on it isn't misread as a real
    Prof/Lab/Total triple.
    """
    results: dict[str, dict[str, float]] = {}
    for line in text.split("\n"):
        code_match = _DD_PDF_LINE_CODE_RE.search(line)
        if not code_match:
            continue
        code = code_match.group(1)
        if code not in known_codes or code in results:
            continue
        numbers = [float(n.replace(",", "")) for n in _DD_PDF_NUMBER_RE.findall(line[code_match.end():])]
        if not numbers:
            continue
        if len(numbers) == 2 and numbers[0] == numbers[1]:
            results[code] = {"prof": numbers[0], "lab": 0.0, "total": numbers[0]}
        elif len(numbers) >= 3:
            prof, lab, total = numbers[0], numbers[1], numbers[-1]
            if abs((prof + lab) - total) < 0.01:
                results[code] = {"prof": prof, "lab": lab, "total": total}
            else:
                results[code] = {"total": numbers[-1]}
        else:
            results[code] = {"total": numbers[-1]}
    return results


def load_pt_dd_fees_from_files(files: list[Path], known_codes: set[str], verbose: bool = True):
    """DD-specific: resolves separate Prof/Lab/Total fees per code (see
    extract_dd_codes_from_rows) from spreadsheet sources whose columns are
    labeled that way, before falling back to the generic single-fee scan
    (load_pt_fees_from_files) for whatever's still missing -- covering
    sources with no distinguishable Prof/Lab/Total columns (plain text/PDF
    guides, docx) as well as spreadsheets without them. A fallback match is
    recorded as {'total': fee}: the generic scanner takes the rightmost
    numeric candidate in a row, which in every DD source seen so far is the
    Total column when one exists, so that's the safest single role to
    attribute it to.
    Returns (dict[code, {'prof': .., 'lab': .., 'total': ..}], sources_used).
    """
    role_fees: dict[str, dict[str, float]] = {}
    sources_used: list[tuple[str, int]] = []

    def _is_regional_variant(f: Path) -> bool:
        return "NORTHERN" in f.stem.upper()

    spreadsheets = sorted(
        (f for f in files if f.suffix.lower() in SPREADSHEET_SUFFIXES),
        key=lambda f: (f.suffix.lower() != ".xlsx", _is_regional_variant(f)),
    )
    for f in spreadsheets:
        remaining = known_codes - role_fees.keys()
        if not remaining:
            break
        try:
            new_fees = extract_dd_codes_from_rows(tables_from_spreadsheet(f), remaining)
        except Exception as e:
            if verbose:
                print(f"    WARNING: failed to read {f.name}: {e}")
            continue
        if new_fees:
            role_fees.update(new_fees)
            sources_used.append((f.name, len(new_fees)))

    # Some docx guides (e.g. AB's) are a long series of small per-section
    # tables, most with no header row of their own, but a handful do repeat
    # a "DAC CODE / PROFESSIONAL FEE / LAB FEE / TOTAL FEE" header -- those
    # sections' Prof/Lab/Total can still be resolved from labels the same
    # way as a spreadsheet's.
    for f in (f for f in files if f.suffix.lower() in DOC_SUFFIXES):
        remaining = known_codes - role_fees.keys()
        if not remaining:
            break
        try:
            new_fees = extract_dd_codes_from_rows(tables_from_docx(f), remaining)
        except Exception as e:
            if verbose:
                print(f"    WARNING: failed to read {f.name}: {e}")
            continue
        if new_fees:
            role_fees.update(new_fees)
            sources_used.append((f.name, len(new_fees)))

    # PDFs whose rows carry Prof/Lab/Total on the same line as the code
    # (see extract_dd_codes_from_pdf_text) -- tried before the generic
    # single-fee fallback below so a PDF source doesn't lose its Prof/Lab
    # breakdown down to Total-only just because it isn't a spreadsheet/docx.
    for f in (f for f in files if f.suffix.lower() in PDF_SUFFIXES):
        remaining = known_codes - role_fees.keys()
        if not remaining:
            break
        try:
            reader = pypdf.PdfReader(str(f))
            layout_text = "\n".join(
                page.extract_text(extraction_mode="layout") or "" for page in reader.pages
            )
            new_fees = extract_dd_codes_from_pdf_text(layout_text, remaining)
        except Exception as e:
            if verbose:
                print(f"    WARNING: failed to read {f.name}: {e}")
            continue
        if new_fees:
            role_fees.update(new_fees)
            sources_used.append((f"{f.name} (pdf rows)", len(new_fees)))

    missing = known_codes - role_fees.keys()
    if missing:
        single_fees, single_sources = load_pt_fees_from_files(files, missing, verbose)
        for code, fee in single_fees.items():
            role_fees[code] = {"total": fee}
        sources_used.extend(single_sources)

    return role_fees, sources_used


def load_pt_dd_fees(specialty_dir: Path, province: str, known_codes: set[str], verbose: bool = True):
    """Resolve DD PT fees for a province from whatever files are available.
    Returns (dict[code, {'prof': .., 'lab': .., 'total': ..}], sources_used, files found)."""
    files = discover_pt_files(specialty_dir, province)
    fees, sources_used = load_pt_dd_fees_from_files(files, known_codes, verbose)
    return fees, sources_used, files


def resolve_dd_role_values(values: dict[str, float]) -> tuple[float | None, float | None, float | None]:
    """Reduce a per-code {'prof', 'lab', 'total'} dict (see
    load_pt_dd_fees) to the (prof, lab, total) triple written to the DD
    sheet. When both Prof and Lab were found, Total is *recomputed* as
    their sum rather than trusted as extracted, even if a labeled Total
    column was also found: some guides' own Total column includes a small
    escalation/adjustment (seen in ON's and NS's sources, off by a couple
    percent) that the reference doesn't carry through, while every
    confirmed-correct source's labeled Total already equals Prof + Lab
    exactly anyway -- so recomputing it is a no-op for those and a fix for
    the rest. Total falls back to whatever was extracted only when Prof or
    Lab is missing, and -- for a source that doesn't distinguish Prof from
    Lab at all (just a single fallback 'total' value) -- that value is
    treated as the Prof fee too, matching this project's original
    single-value DD behavior for sources without labeled columns.
    """
    prof = values.get("prof")
    lab = values.get("lab")
    total = values.get("total")
    if prof is not None and lab is not None:
        total = prof + lab
    elif total is None:
        if prof is not None:
            total = prof
        elif lab is not None:
            total = lab
    if prof is None and lab is None and total is not None:
        prof = total
    return prof, lab, total


# Maps a CDCP SP sub-specialty code to name fragments that identify a PT fee
# guide file as specific to that sub-specialty (e.g. "ON PA Fee Guide.xlsx",
# "MDA 2026 Periodontics....xlsx" both indicate Periodontics -> "PA"). Used
# to prefer a sub-specialty-specific guide's fee over a general/all-specialty
# guide's fee for the same code, since the same procedure code commonly has
# a genuinely different fee depending on which specialty bills it.
SUBSPECIALTY_FILE_MARKERS: dict[str, list[str]] = {
    "EN": ["EN", "END", "ENDO", "ENDODONTIC", "ENDODONTICS"],
    "OS": ["OS", "OMS", "ORAL SURGERY", "ORAL AND MAXILLOFACIAL SURGERY", "MAXILLOFACIAL"],
    "PA": ["PA", "PER", "PERIODONTIC", "PERIODONTICS", "PERIODONTOLOGY"],
    "PE": ["PE", "PED", "PEDIATRIC", "PEDIATRICS", "PAEDIATRIC", "PAEDIATRICS"],
    "PR": ["PR", "PROSTHODONTIC", "PROSTHODONTICS"],
    "OM": ["OM", "ORAL MEDICINE"],
    "OP": ["OP", "ORAL PATHOLOGY"],
    "OR": ["OR", "ORT", "ORTHODONTIC", "ORTHODONTICS"],
    "RA": ["RA", "RADIOLOGY"],
    "AN": ["AN", "ANESTHESIA", "ANESTHESIOLOGY"],
}


# Reverse lookup from a specialty *label* (as it might appear in a source's
# own per-row specialty column, e.g. "END", "OMS", "PER") to the CDCP
# sub-specialty code it means -- reuses the same marker vocabulary as
# filename classification (see classify_file_specialty), since combined
# guides tend to abbreviate specialties the same way whether in a filename
# or a column value. "GP" and "LTC" aren't CDCP sub-specialties but share
# this column in combined guides, so they get their own pseudo-entries --
# purely so a GP/LTC row is recognized as *labeled* (and therefore excluded
# from matching any real sub-specialty, and from the "unlabeled" fallback
# bucket -- an LTC-context rate isn't the general rate for whichever
# sub-specialty the code also happens to belong to) rather than being
# mistaken for an unlabeled row (see find_row_specialty_column).
_ALL_SPECIALTY_MARKERS: dict[str, str] = {"GP": "GP", "LTC": "LTC"}
for _code, _markers in SUBSPECIALTY_FILE_MARKERS.items():
    for _marker in _markers:
        _ALL_SPECIALTY_MARKERS.setdefault(_marker, _code)


def _classify_specialty_cell(cell) -> str | None:
    if not isinstance(cell, str):
        return None
    return _ALL_SPECIALTY_MARKERS.get(cell.strip().upper())


def find_row_specialty_column(rows) -> int | None:
    """Find a column whose values are recognizable CDCP sub-specialty labels
    (e.g. "END", "OMS", "PER", "GP") rather than a specialty-specific
    *file*. Some combined SP guides (e.g. PE's) list every sub-specialty's
    codes together in one file with a per-row specialty column instead of
    splitting into separate files/sections -- without identifying that
    column, the same code appearing under two different specialties (e.g.
    "25781" priced differently under "GP" and under "END") can't be told
    apart, and whichever row happens to come first wins regardless of which
    specialty was actually asked for.

    Returns the column index with the most matches among a sample of rows,
    if at least a meaningful fraction of that sample matches -- a low
    fraction means this source likely doesn't have a real specialty column
    at all (a stray "OR" or "PA" abbreviation elsewhere shouldn't count).
    """
    sample = rows[:500]
    if not sample:
        return None
    width = max((len(r) for r in sample), default=0)
    best_idx, best_count = None, 0
    for i in range(width):
        count = sum(1 for r in sample if i < len(r) and _classify_specialty_cell(r[i]) is not None)
        if count > best_count:
            best_idx, best_count = i, count
    if best_idx is not None and best_count >= len(sample) * 0.3:
        return best_idx
    return None


def classify_file_specialty(path: Path, province: str | None = None) -> set[str]:
    """Which CDCP sub-specialty code(s), if any, a PT fee guide filename
    identifies (e.g. "ON PA Fee Guide 2026.xlsx" -> {"PA"}). Empty set means
    the file isn't specific to one sub-specialty (e.g. a general/combined
    guide like "ON DA Fee Guide" or "BC LTC Fee Guide").

    If `province` is given, its aliases (see PROVINCE_ALIASES) are stripped
    from the start of the name before matching -- otherwise a filename like
    "PE GP SP LTC Fee Guide" (Prince Edward Island's combined guide) gets
    misread as Pediatric-specific, since "PE" is coincidentally both the
    province's abbreviation and the Pediatric specialty marker.
    """
    name = path.stem.upper()
    if province:
        for alias in PROVINCE_ALIASES.get(province, [province]):
            m = re.match(rf"^{re.escape(alias)}\b\s*", name, re.IGNORECASE)
            if m:
                name = name[m.end():]
                break
    matches = set()
    for code, markers in SUBSPECIALTY_FILE_MARKERS.items():
        for marker in markers:
            if re.search(rf"\b{re.escape(marker)}\b", name):
                matches.add(code)
                break
    return matches


# Filename markers for guides that are scoped to a specific care *setting*
# (e.g. Long Term Care) rather than a specific CDCP sub-specialty. These
# aren't a real CDCP specialty code, so a file scoped *only* to one of these
# settings must not be treated as a general/blanket fallback for every
# sub-specialty that lacks its own guide (a code's LTC-context rate isn't
# its regular-context rate). But some provinces publish one combined guide
# covering GP, SP, *and* LTC together (e.g. "PE GP SP LTC Fee Guide") --
# that file is a legitimate general SP source despite mentioning LTC, since
# it isn't LTC-exclusive.
_CONTEXT_RESTRICTED_MARKERS = ["LTC"]
_COMBINED_GUIDE_MARKERS = ["GP", "SP"]


def _is_context_restricted(path: Path) -> bool:
    name = path.stem.upper()
    has_restricted_marker = any(re.search(rf"\b{marker}\b", name) for marker in _CONTEXT_RESTRICTED_MARKERS)
    if not has_restricted_marker:
        return False
    is_combined_guide = any(re.search(rf"\b{marker}\b", name) for marker in _COMBINED_GUIDE_MARKERS)
    return not is_combined_guide


# Some guides (so far only PE's combined GP+SP PDF) define a specialist's
# fee for an entire numeric code range as a flat percentage markup over the
# general practitioner's fee, rather than listing individual specialist
# fees at all -- e.g. Appendix F: "SERVICES PROVIDED BY A PROSTHODONTIST /
# SECTION 50000 - 59999 / FEES FOR ALL CODES 20% HIGHER THAN FOR GENERAL
# PRACTITIONER'S SUGGESTED FEE." Used as a per-code fallback in
# load_pt_fees_by_subspecialty for codes in that range with no
# specialist-specific rate elsewhere.
_SPECIALIST_NOUN_TO_SUBSPECIALTY: dict[str, str] = {
    "PROSTHODONTIST": "PR",
    "PERIODONTIST": "PA",
    "ENDODONTIST": "EN",
    "ORTHODONTIST": "OR",
    "ORAL AND MAXILLO-FACIAL SURGEON": "OS",
    "ORAL SURGEON": "OS",
    "PAEDIATRIC DENTIST": "PE",
    "PEDIATRIC DENTIST": "PE",
}
_MULTIPLIER_APPENDIX_RE = re.compile(
    r"SERVICES PROVIDED BY (?:A |AN |CERTIFIED )*([A-Za-z][A-Za-z \-]*?)\s*\n"
    r"SECTION\s*(\d+)\s*-\s*(\d+)\s*\n"
    r"FEES FOR ALL CODES\s*(\d+)\s*%\s*HIGHER THAN FOR GENERAL PRACTITIONER",
    re.IGNORECASE,
)


def find_specialist_multiplier_ranges(text: str) -> list[tuple[str, int, int, float]]:
    """Returns a list of (sub_specialty_code, range_low, range_high,
    multiplier) parsed from "FEES FOR ALL CODES N% HIGHER..." appendix
    blocks in `text` (see comment above)."""
    results = []
    for m in _MULTIPLIER_APPENDIX_RE.finditer(text):
        noun = re.sub(r"\s+", " ", m.group(1)).strip().upper()
        code = _SPECIALIST_NOUN_TO_SUBSPECIALTY.get(noun)
        if code is None:
            continue
        low, high, pct = int(m.group(2)), int(m.group(3)), int(m.group(4))
        results.append((code, low, high, 1 + pct / 100))
    return results


def _multiplier_for_code(code: str, ranges: list[tuple[int, int, float]]) -> float | None:
    try:
        n = int(code)
    except ValueError:
        return None
    for low, high, mult in ranges:
        if low <= n <= high:
            return mult
    return None


def load_pt_fees_by_subspecialty(
    specialty_dir: Path,
    province: str,
    codes_by_subspecialty: dict[str, set[str]],
    gp_specialty_dir: Path | None = None,
    verbose: bool = True,
):
    """SP-style resolution: like load_pt_fees, but aware that the same code
    can have a different fee under different sub-specialties. Files whose
    name identifies a specific sub-specialty (see classify_file_specialty)
    are tried first for that sub-specialty's codes, before falling back to
    general/unspecific files (which is all load_pt_fees does on its own).

    If `gp_specialty_dir` is given and a sub-specialty gets *no* matches at
    all from SP sources (i.e. no PT guide covers that specialty for this
    province -- e.g. BC has no EN/OM/OP/OS/PR/RA-specific guide), its codes
    fall back to the province's GP fee guide instead -- the reference file's
    own documented convention ("For any SP code without SP specific fee, GP
    fee is assumed"). This is deliberately an all-or-nothing trigger per
    sub-specialty, not a per-missing-code one: a sub-specialty with mostly
    good SP-specific coverage and a few individually-unmatched codes is more
    likely suffering an extraction gap in its own guide than a genuine
    absence of specialty-specific pricing, and guessing the GP fee for those
    stray gaps does more harm (contaminating otherwise-correct data) than
    leaving them "N/A".

    Before that blanket fallback, though, any codes covered by a documented
    percentage-markup rule (see find_specialist_multiplier_ranges -- so far
    only PE's guide) get GP fee x that specialty's stated multiplier
    instead of the plain GP fee, since that's a known, precise rule rather
    than a guess.

    Returns (fees dict keyed by (code, sub_specialty), sources_used, files found).
    """
    files = discover_pt_files(specialty_dir, province)
    general_files = [f for f in files
                      if not classify_file_specialty(f, province) and not _is_context_restricted(f)]

    multiplier_ranges: dict[str, list[tuple[int, int, float]]] = {}
    for f in files:
        if f.suffix.lower() != ".pdf":
            continue
        try:
            reader = pypdf.PdfReader(str(f))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception:
            continue
        for code, low, high, mult in find_specialist_multiplier_ranges(text):
            multiplier_ranges.setdefault(code, []).append((low, high, mult))

    fees: dict[tuple[str, str], float] = {}
    sources_used: list[tuple[str, int]] = []
    for sub_specialty, codes in codes_by_subspecialty.items():
        specific_files = [f for f in files if sub_specialty in classify_file_specialty(f, province)]
        candidate_files = specific_files + general_files if specific_files else general_files
        sub_fees, sub_sources = load_pt_fees_from_files(
            candidate_files, codes, verbose, target_specialty=sub_specialty
        )
        for code, fee in sub_fees.items():
            fees[(code, sub_specialty)] = fee
        sources_used.extend(sub_sources)

        ranges = multiplier_ranges.get(sub_specialty)
        if ranges and gp_specialty_dir is not None:
            # Applied to every code in range, not just ones sub_fees missed
            # -- a documented multiplier rule overrides even a value
            # sub_fees *did* find, since that match is usually the generic
            # PDF reader (not specialty-aware, unlike the spreadsheet/csv/
            # docx readers -- see extract_codes_from_rows) grabbing the
            # wrong, non-specialist section of the same combined document
            # rather than genuinely finding this specialty's own rate.
            in_range = {c for c in codes if _multiplier_for_code(c, ranges) is not None}
            if in_range:
                gp_fees, _, _ = load_pt_fees(gp_specialty_dir, province, in_range, verbose=False)
                applied = 0
                for code in in_range:
                    gp_fee = gp_fees.get(code)
                    mult = _multiplier_for_code(code, ranges)
                    if gp_fee is None or mult is None:
                        continue
                    fees[(code, sub_specialty)] = gp_fee * mult
                    applied += 1
                if applied:
                    sources_used.append((f"GP fee x specialist markup ({applied})", applied))

        if not sub_fees and gp_specialty_dir is not None:
            gp_fees, gp_sources, _ = load_pt_fees(gp_specialty_dir, province, codes, verbose=False)
            for code, fee in gp_fees.items():
                if (code, sub_specialty) not in fees:
                    fees[(code, sub_specialty)] = fee
            if gp_fees:
                sources_used.extend((f"{label} (as GP fallback)", n) for label, n in gp_sources)

    return fees, sources_used, files
